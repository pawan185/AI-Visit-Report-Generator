import re
from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document(report_text: str) -> BytesIO:
    """
    Takes the combined report text (which contains markdown) and generates 
    a Word Document in memory, returning a BytesIO object for Streamlit.
    """
    doc = Document()
    doc.add_heading('Final Visit Report & Proposal', 0)
    
    for line in report_text.split('\n'):
        line = line.strip()
        if line == "":
            continue
            
        # Parse Page Break / Section Divider
        if line.startswith('---'):
            doc.add_page_break()
            continue
            
        # Parse Markdown Headers
        if line.startswith('### '):
            p = doc.add_heading(line.replace('###', '').strip(), level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line.replace('##', '').strip(), level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line.replace('#', '').strip(), level=1)
        # Parse Bullets vs Normal Paragraphs
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            line = line.lstrip("-* ")
        else:
            p = doc.add_paragraph()
            
        # Professional Formatting: Justify text and set ultra-tight space between paragraphs
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        
        # Robustly parse inline bold formatting (e.g. **text**)
        if not (line.startswith('#') or line.startswith('##') or line.startswith('###')):
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for i, part in enumerate(parts):
                # Clean up stray single asterisks
                clean_part = part.replace('*', '')
                if not clean_part:
                    continue
                    
                run = p.add_run(clean_part)
                if i % 2 == 1:
                    run.bold = True
                
    # Save the document to an in-memory byte stream
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
